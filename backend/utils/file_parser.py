# when running use this command in the local :
import os
import io
import re
import logging
import zipfile

from typing import List
from xml.etree import ElementTree as ET

from PyPDF2 import PdfReader
from docx import Document

logger = logging.getLogger(__name__)

WORD_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
DOCX_PAGE_NUMBER_PATTERN = re.compile(
    r"^(?:page\s*)?\d+(?:\s*of\s*\d+)?$", re.IGNORECASE
)

# ----------------------------
# Optional PDF extraction deps
# ----------------------------
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    from pdf2image import convert_from_bytes
    import pytesseract
    HAS_OCR = True
except ImportError:
    HAS_OCR = False


# ─────────────────────────────────────────────────────────────────────────────
# FIX #11: Whitespace normalisation
# ─────────────────────────────────────────────────────────────────────────────

def normalize_whitespace(text: str) -> str:
    """
    FIX #11: Aggressively normalise extracted text whitespace.

    Operations (in order):
      1. Replace tabs with a single space.
      2. Replace non-breaking spaces (\\xa0) with regular spaces.
      3. Collapse multiple consecutive spaces within a line to one space.
      4. Strip leading/trailing spaces from every line.
      5. Collapse runs of 3+ consecutive blank lines to 2 blank lines
         (preserving paragraph structure).

    This prevents issues like:
        "John    Doe        Senior    Engineer"
        → "John Doe Senior Engineer"
    """
    if not text:
        return text

    # Step 1 & 2: tabs and non-breaking spaces → regular space
    text = text.replace('\t', ' ').replace('\xa0', ' ')

    # Step 3 & 4: collapse spaces and strip each line
    lines = text.split('\n')
    lines = [re.sub(r' +', ' ', line).strip() for line in lines]
    text = '\n'.join(lines)

    # Step 5: at most two consecutive blank lines (paragraph break)
    text = re.sub(r'\n{3,}', '\n\n', text)

    return text


# ----------------------------
# PDF extractors
# ----------------------------

def _extract_text_from_pdf_with_pdfplumber(file_bytes: bytes) -> str:
    """Extract text using pdfplumber (often best for complex layouts)."""
    if not HAS_PDFPLUMBER:
        return ""

    text_parts: List[str] = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page_idx, page in enumerate(pdf.pages, start=1):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(page_text)
                logger.info(
                    f"📄 pdfplumber page {page_idx}: {len(page_text or '')} chars"
                )
    except Exception as e:
        logger.warning(f"pdfplumber extraction failed: {e}")
        return ""

    return "\n".join(text_parts).strip()


def _extract_text_from_pdf_with_pypdf2(file_bytes: bytes) -> str:
    """Extract text using PyPDF2 as fallback."""
    text_parts: List[str] = []
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        for page_idx, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_parts.append(page_text)
            logger.info(
                f"📄 PyPDF2 page {page_idx}: {len(page_text)} chars"
            )
    except Exception as e:
        logger.warning(f"PyPDF2 extraction failed: {e}")
        return ""

    return "\n".join(text_parts).strip()


def _extract_text_from_pdf_with_ocr(file_bytes: bytes) -> str:
    """Extract text using OCR (for scanned/image-based PDFs)."""
    if not HAS_OCR:
        logger.warning("OCR not available (missing pdf2image and/or pytesseract).")
        return ""

    text_parts: List[str] = []
    try:
        images = convert_from_bytes(file_bytes, dpi=300)
        for i, image in enumerate(images, start=1):
            try:
                page_text = pytesseract.image_to_string(image, lang="eng") or ""
                if page_text.strip():
                    text_parts.append(page_text.strip())
                logger.info(f"🧾 OCR page {i}: {len(page_text)} chars")
            except Exception as e:
                logger.warning(f"OCR failed for page {i}: {e}")
                continue
    except Exception as e:
        logger.warning(f"OCR extraction failed: {e}")
        return ""

    return "\n".join(text_parts).strip()


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text from PDF using multiple methods with fallback:
    1) pdfplumber
    2) PyPDF2
    3) OCR (if available)
    """
    text = _extract_text_from_pdf_with_pdfplumber(file_bytes)
    if text:
        logger.info("✅ PDF extracted successfully using pdfplumber")
        return text

    text = _extract_text_from_pdf_with_pypdf2(file_bytes)
    if text:
        logger.info("✅ PDF extracted successfully using PyPDF2")
        return text

    logger.info("⚠️ Standard PDF extraction failed; attempting OCR...")
    text = _extract_text_from_pdf_with_ocr(file_bytes)
    if text:
        logger.info("✅ PDF extracted successfully using OCR")
        return text

    logger.warning("❌ All PDF extraction methods failed")
    return ""


# ----------------------------
# DOCX extractor
# ----------------------------

def _normalize_docx_line(text: str) -> str:
    """Normalize whitespace while preserving visible text content."""
    normalized = (text or "").replace("\xa0", " ")
    normalized = re.sub(r"\s+", " ", normalized).strip()
    return normalized


def _is_layout_noise_line(text: str) -> bool:
    """Filter common page-number artifacts from headers/footers."""
    return bool(DOCX_PAGE_NUMBER_PATTERN.fullmatch((text or "").strip()))


def _extract_lines_from_docx_xml_part(xml_bytes: bytes) -> List[str]:
    """
    Extract visible paragraph lines from a DOCX XML part.
    Captures text in normal paragraphs, tables, and text boxes.
    """
    root = ET.fromstring(xml_bytes)
    lines: List[str] = []

    for paragraph in root.findall(".//w:p", WORD_NS):
        chunks: List[str] = []
        for node in paragraph.iter():
            tag = node.tag.rsplit("}", 1)[-1] if "}" in node.tag else node.tag

            if tag == "t" and node.text:
                chunks.append(node.text)
            elif tag == "tab":
                chunks.append(" ")
            elif tag in {"br", "cr"}:
                chunks.append("\n")

        joined = "".join(chunks)
        for line in joined.split("\n"):
            normalized = _normalize_docx_line(line)
            if normalized and not _is_layout_noise_line(normalized):
                lines.append(normalized)

    return lines


def _extract_text_from_docx_xml(file_bytes: bytes) -> str:
    """
    Extract DOCX text from XML parts (headers + body + footers).
    This recovers text frequently missed by python-docx (e.g., text boxes).
    """
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes), "r") as docx_zip:
            part_names = set(docx_zip.namelist())

            headers = sorted(
                name for name in part_names
                if name.startswith("word/header") and name.endswith(".xml")
            )
            footers = sorted(
                name for name in part_names
                if name.startswith("word/footer") and name.endswith(".xml")
            )

            ordered_parts: List[str] = []
            ordered_parts.extend(headers)
            if "word/document.xml" in part_names:
                ordered_parts.append("word/document.xml")
            ordered_parts.extend(footers)

            all_lines: List[str] = []
            for part_name in ordered_parts:
                all_lines.extend(
                    _extract_lines_from_docx_xml_part(docx_zip.read(part_name))
                )

            # Remove only immediate duplicates (common with repeated line wraps).
            deduped_lines: List[str] = []
            previous_key = None
            for line in all_lines:
                key = line.casefold()
                if key == previous_key:
                    continue
                deduped_lines.append(line)
                previous_key = key

            return "\n".join(deduped_lines).strip()
    except Exception as e:
        logger.warning(f"DOCX XML extraction failed: {e}")
        return ""


def _extract_text_from_docx_with_python_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx paragraphs + tables."""
    try:
        doc = Document(io.BytesIO(file_bytes))
        parts: List[str] = []

        for para in doc.paragraphs:
            t = (para.text or "").strip()
            if t:
                parts.append(t)

        for table in doc.tables:
            for row in table.rows:
                row_cells: List[str] = []
                for cell in row.cells:
                    ct = (cell.text or "").strip()
                    if ct:
                        row_cells.append(ct)
                if row_cells:
                    parts.append(" | ".join(row_cells))

        return "\n".join(parts).strip()
    except Exception as e:
        logger.warning(f"DOCX python-docx extraction failed: {e}")
        return ""


def _extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from DOCX using XML-aware extraction first, then fallback.
    """
    text = _extract_text_from_docx_xml(file_bytes)
    if text:
        logger.info("✅ DOCX extracted successfully using XML parser")
        return text

    text = _extract_text_from_docx_with_python_docx(file_bytes)
    if text:
        logger.info("✅ DOCX extracted successfully using python-docx")
        return text

    logger.warning("❌ All DOCX extraction methods failed")
    return ""


# ----------------------------
# TXT extractor
# ----------------------------

def _extract_text_from_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore").strip()


# ─────────────────────────────────────────────────────────────────────────────
# PUBLIC API (keep signature identical)
# ─────────────────────────────────────────────────────────────────────────────

def extract_text_from_file(file_path: str) -> str:
    """
    Extract text from an uploaded file using direct parsing (path-based API).

    FIX #11: normalize_whitespace() is applied to all extracted text before
    returning, ensuring no multi-space runs, tabs, or excessive blank lines
    reach the chunking / LLM pipeline.

    Args:
        file_path: Path to the uploaded file

    Returns:
        Extracted and whitespace-normalised text content
    """
    file_extension = os.path.splitext(file_path)[1].lower()
    logger.info(f"🔍 Processing file: {file_path} (extension: {file_extension})")

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    file_size = os.path.getsize(file_path)
    logger.info(f"📊 File size: {file_size} bytes")

    try:
        with open(file_path, "rb") as f:
            file_bytes = f.read()

        if not file_bytes:
            logger.error("❌ File content is empty")
            raise ValueError("File is empty")

        if file_extension == ".pdf":
            logger.info("🔄 Extracting PDF (pdfplumber → PyPDF2 → OCR fallback)...")
            text = _extract_text_from_pdf(file_bytes)

        elif file_extension == ".docx":
            logger.info("🔄 Extracting DOCX (XML-aware parser → python-docx fallback)...")
            text = _extract_text_from_docx(file_bytes)

        elif file_extension == ".txt":
            logger.info("🔄 Reading TXT...")
            text = _extract_text_from_txt(file_bytes)

        else:
            raise ValueError(f"Unsupported file type: {file_extension}")

        # ✅ FIX #11: Normalise whitespace before handing to pipeline
        text = normalize_whitespace(text)

        logger.info(f"✅ Extraction + normalisation successful – {len(text)} characters")
        logger.info(f"📝 First 500 chars: {text[:500]}...")
        return text

    except Exception as e:
        logger.error(f"❌ File extraction failed: {e}")
        raise Exception(f"Failed to extract text from file: {e}")